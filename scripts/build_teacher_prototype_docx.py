from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "AI翻译App原型设计交付文档.md"
OUTPUT = ROOT / "docs" / "AI翻译App原型设计交付文档.docx"


def set_run_font(run, font_name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_style_font(style, font_name="Microsoft YaHei", size=11, color="222222", bold=None):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, bottom=90, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def clean_inline(text):
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return text.strip()


def add_rich_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part.strip("`*"))
        is_strong = part.startswith("**") and part.endswith("**")
        is_code = part.startswith("`") and part.endswith("`")
        set_run_font(run, "Microsoft YaHei", 10.5, "222222", bold=is_strong)
        if is_code:
            set_run_font(run, "Consolas", 10, "333333", bold=False)
    return p


def add_markdown_table(doc, rows):
    headers = [clean_inline(c) for c in rows[0]]
    body = [[clean_inline(c) for c in row] for row in rows[1:]]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    if headers == ["成员", "核心任务", "占比", "责任承诺"]:
        widths = [900, 2500, 760, 5200]
    else:
        widths = [int(9360 / len(headers))] * len(headers)
    set_table_width(table, widths)

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "4472C4")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, "Microsoft YaHei", 10.5, "FFFFFF", True)

    for row_data in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, "Microsoft YaHei", 10, "222222", False)

    doc.add_paragraph()


def add_markdown_image(doc, alt_text, image_path):
    resolved = (SOURCE.parent / image_path).resolve()
    if not resolved.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[原型图缺失：{alt_text}]")
        set_run_font(run, "Microsoft YaHei", 10, "9B1C1C", True)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # 手机长图按单页可容纳的宽度插入，兼顾细节可读性和页面不溢出。
    run.add_picture(str(resolved), width=Inches(2.9))


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    set_style_font(styles["Normal"], size=10.5, color="222222")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    set_style_font(styles["Title"], size=22, color="1F3A5F", bold=True)
    styles["Title"].paragraph_format.space_after = Pt(12)
    set_style_font(styles["Heading 1"], size=16, color="2E74B5", bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    set_style_font(styles["Heading 2"], size=13, color="2E74B5", bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 3"], size=12, color="1F4D78", bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    set_style_font(styles["List Bullet"], size=10.5, color="222222")
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Bullet"].paragraph_format.space_after = Pt(4)
    set_style_font(styles["List Number"], size=10.5, color="222222")
    styles["List Number"].paragraph_format.left_indent = Inches(0.5)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.25)
    styles["List Number"].paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "AI 翻译 App 原型设计交付文档"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, "Microsoft YaHei", 9, "666666")

    footer = section.footer.paragraphs[0]
    footer.text = "原型设计说明 · 翻译闭环校验 · 成员责任承诺"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, "Microsoft YaHei", 9, "888888")

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_buffer = []
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    def flush_code():
        nonlocal code_buffer
        if code_buffer:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run("\n".join(code_buffer))
            set_run_font(run, "Consolas", 10, "333333")
            code_buffer = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            flush_table()
            if in_code:
                in_code = False
                flush_code()
            else:
                in_code = True
                code_buffer = []
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
                continue
            table_rows.append(cells)
            continue

        flush_table()

        if not line.strip():
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            set_run_font(run, "Microsoft YaHei", 22, "1F3A5F", True)
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle.add_run("原型设计说明 · 翻译闭环校验 · 成员责任承诺")
            set_run_font(run, "Microsoft YaHei", 11, "666666")
            continue

        if line.startswith("## "):
            doc.add_paragraph(clean_inline(line[3:]), style="Heading 1")
            continue

        if line.startswith("### "):
            doc.add_paragraph(clean_inline(line[4:]), style="Heading 2")
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
        if image_match:
            add_markdown_image(doc, image_match.group(1), image_match.group(2))
            continue

        if line.startswith("- [x] "):
            p = add_rich_paragraph(doc, "已完成：" + line[6:].strip(), style="List Bullet")
            continue

        if line.startswith("- "):
            add_rich_paragraph(doc, line[2:].strip(), style="List Bullet")
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        if numbered:
            add_rich_paragraph(doc, numbered.group(2), style="List Number")
            continue

        add_rich_paragraph(doc, line)

    flush_table()
    flush_code()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
